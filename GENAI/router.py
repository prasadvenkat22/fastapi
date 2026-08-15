from fastapi import APIRouter, HTTPException, File, UploadFile, Form
from typing import List, Optional
import io
import os
import uuid
import httpx
from pydantic import BaseModel

from .llm_integration import DocumentRequest, LLMRequest, LLMResponse, DocumentResponse
from .rag_service import generate_response, run_rag
from .vector_stores import vector_store_factory, VectorDocument
from .agents.supervisor import run_supervisor

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

router = APIRouter(prefix="/api/genai", tags=["GENAI"])


class QueryResponse(BaseModel):
    results: List[DocumentResponse]


@router.post("/query", response_model=QueryResponse)
async def genai_query(request: DocumentRequest, vector_store_name: str = "pgvector", top_k: int = 2):
    """Accept a structured request and return structured results from a vector store."""
    try:
        vs = vector_store_factory(vector_store_name, request.embedding_provider)
        results = await vs.get_documents(query=request.query, top_k=request.top_k, filters=request.filters)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    return QueryResponse(results=results)


@router.post("/llm", response_model=LLMResponse)
async def genai_llm(request: LLMRequest):
    """Accept a structured LLM request and return a structured LLM response."""
    try:
        response = await generate_response(request)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return response


def _process_file(data: bytes, filename: str) -> str:
    text = ""
    lower_filename = filename.lower()

    if lower_filename.endswith(".pdf") and PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(data))
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n\n".join(pages)
        except Exception:
            text = ""
    elif lower_filename.endswith((".txt", ".csv")):
        try:
            text = data.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
    elif lower_filename.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            text = "\n\n".join([p.text for p in doc.paragraphs])
        except Exception:
            text = ""
    elif any(lower_filename.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
        try:
            from PIL import Image
            import pytesseract
            img = Image.open(io.BytesIO(data))
            text = pytesseract.image_to_string(img)
        except Exception:
            text = "[binary image content omitted]"
    else:
        try:
            text = data.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
    return text


async def _register_or_update_user(username: str, email: Optional[str] = None) -> bool:
    """
    Register or update user in MongoDB via the /api/mongo/users/ endpoint.
    Returns True if successful, False otherwise.
    """
    try:
        # Try to create user in MongoDB
        # Generate a valid email from username
        valid_email = email if email else f"{username.replace('_', '').replace(' ', '')}@genai.app"

        user_data = {
            "name": username[:16],  # Ensure name is within max length
            "email": valid_email,
            "password": "GenAI@2024",  # Default password (8-15 chars)
            "tenantdb": "genai",
            "application": "rag",
            "role": "user",
            "status": True
        }

        # Use internal API call to register user
        base_url = os.getenv("API_BASE_URL", "http://localhost:8000")

        async with httpx.AsyncClient(timeout=10.0) as client:
            # Check if user already exists
            try:
                response = await client.get(f"{base_url}/api/mongo/users/")
                if response.status_code == 200:
                    existing_users = response.json()
                    # Check if username already exists
                    if any(u.get("name") == username for u in existing_users):
                        return True  # User already exists, no need to create
            except Exception:
                pass  # If check fails, proceed to create

            # Create new user
            response = await client.post(
                f"{base_url}/api/mongo/users/",
                json=user_data
            )

            return response.status_code in [200, 201]
    except Exception as e:
        # Log error but don't fail the upload
        print(f"Warning: Could not register user '{username}': {e}")
        return False


@router.post("/query/upload", response_model=QueryResponse)
async def genai_query_upload(
    files: List[UploadFile] = File(...),
    query: str = Form(...),
    username: str = Form("test_user"),
    vector_store_name: str = Form("pgvector"),
    embedding_provider: str = Form("voyage"),
    llm_provider: str = Form("anthropic"),
    llm_model: str = Form("claude-opus-5"),
    max_tokens: int = Form(512),
    top_k: int = Form(2),
):
    """
    Accept file uploads, add them to the vector store, and then run a RAG query.

    Parameters:
    - files: List of files to upload (PDF, CSV, TXT, DOCX, images)
    - query: Question to ask about the uploaded documents
    - username: User who uploaded the file (default: test_user)
    - vector_store_name: Vector store to use (default: pgvector)
    - embedding_provider: Embedding provider (default: voyage)
    - llm_provider: LLM provider (default: anthropic)
    - llm_model: LLM model (default: claude-opus-5)
    - max_tokens: Maximum tokens for LLM response (default: 800)
    - top_k: Number of similar documents to retrieve (default: 4)
    """
    import datetime

    documents: List[VectorDocument] = []
    for f in files:
        data = await f.read()
        filename = getattr(f, "filename", "uploaded_file")
        content = _process_file(data, filename)
        if content:
            doc_id = str(uuid.uuid4())
            metadata = {
                "filename": filename,
                "username": username,
                "uploaded_at": datetime.datetime.utcnow().isoformat(),
                "file_size": len(data)
            }
            documents.append(VectorDocument(id=doc_id, text=content, metadata=metadata))
        await f.close()

    # Register or update user in MongoDB
    await _register_or_update_user(username)

    try:
        vs = vector_store_factory(vector_store_name, embedding_provider)
        if documents:
            await vs.add_documents(documents)

        doc_request = DocumentRequest(query=query, top_k=top_k, embedding_provider=embedding_provider)
        
        # The uploads parameter is no longer needed since the documents are in the vector store
        results = await run_rag(doc_request, vs, uploads=None)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    return QueryResponse(results=results)


class AgentUploadResponse(BaseModel):
    final_answer: str
    csv_answer: Optional[str] = None
    pdf_answer: Optional[str] = None
    agents_used: List[str]


@router.post("/agent/upload", response_model=AgentUploadResponse)
async def genai_agent_upload(
    files: List[UploadFile] = File(...),
    query: str = Form(...),
):
    """
    Multi-step LangGraph supervisor: uploads one or more CSV/PDF files, and the
    supervisor routes the query to a pandas-dataframe agent (CSV), a RAG Q&A
    agent (PDF), or both — synthesizing a single answer when both are present.

    Parameters:
    - files: List of files to upload (CSV and/or PDF)
    - query: Question to ask about the uploaded file(s)
    """
    csv_text: Optional[str] = None
    pdf_text: Optional[str] = None

    for f in files:
        data = await f.read()
        filename = getattr(f, "filename", "uploaded_file")
        content = _process_file(data, filename)
        if filename.lower().endswith(".csv"):
            csv_text = content
        elif filename.lower().endswith(".pdf"):
            pdf_text = content
        await f.close()

    if not csv_text and not pdf_text:
        raise HTTPException(status_code=400, detail="Upload at least one .csv or .pdf file.")

    try:
        state = run_supervisor(query=query, csv_text=csv_text, pdf_text=pdf_text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    agents_used = []
    if state.get("csv_answer"):
        agents_used.append("csv_agent")
    if state.get("pdf_answer"):
        agents_used.append("pdf_agent")

    return AgentUploadResponse(
        final_answer=state.get("final_answer", ""),
        csv_answer=state.get("csv_answer"),
        pdf_answer=state.get("pdf_answer"),
        agents_used=agents_used,
    )
